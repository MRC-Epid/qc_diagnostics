######################################################################################################
# This script combines the individual results from QC diagnostics and pulls out the variables that should be reviewed to look for monitor faults.
# Version: 1
# Date: 25/10/2024
# Author: cas254
######################################################################################################

# Importing packages
import os
import pandas as pd
import docx
from datetime import date

from Tools.scripts.combinerefs import combine
from docx.shared import RGBColor
import operator
import setup
from setup import NAME_JOB_LIST

##################################################################################


# The settings below refer to the setup settings that were applied when setting up monitors, the expected file length and minimum duration wear.
# As default is it set up to check if device was set up for 7 days at 100 hz, with a sample range of 8g and with a minimum wear of 4 days (96 hours - with at least 6 hours wear in each quadrant).
# These settings can be kept or changed to fit your project.

# --- STUDY SPECIFIC SETTINGS --- #
# These settings can be edited before running the script, the comments explains how they should be changed:
frequency = 100 # What frequency the device was set up to record with (hz)
sample_range = 8 # What sample range the device was set up to record with (g)
min_file_duration = 6.5 # If the device is set up to record for 7 days, this should be set to 6.5 to have a buffer. Change according to duration settings.
max_file_duration = 7.5 # If the device is set up to record for 7 days, this should be set to 7.5 to have a buffer. Change according to duration settings.
min_wear = 96   # Specify how many hours device should be worn as a minimum
quad_wear = 6   # Specify how many hours wear you need for each quadrant for the file to be valid
device_filename = 'No'  # Set this to 'Yes' if you have set OmGui to add device number in the filename when you download the data, set to 'No' if not.

# FROM HERE NOTHING ELSE SHOULD BE EDITED #

# Todays date:
DATE = date.today().strftime("%d%b%Y")

operators = {
    '!=': operator.ne,
    '==': operator.eq,
    '>': operator.gt,
    '<': operator.lt,
    '>=': operator.ge,
    '<=': operator.le
}

################################################################################################################

# Create a file list of all .csv in the results folder
def create_filelist(outputted_results):
    # Change to the directory containing the results
    if not os.path.exists(outputted_results):
        print(f"Directory not found: {outputted_results}")
        return

    os.chdir(outputted_results)
    print(f"Changed directory to: {outputted_results}")

    # Use Python to find all .csv files and write to filelist.txt
    csv_files = [f for f in os.listdir(os.path.join(setup.PROJECT_DIR, setup.RESULTS_FOLDER)) if f.endswith('.csv')]

    if not csv_files:
        print("No CSV files found in the directory.")
        return

    else:
        return csv_files


def read_and_append(csv_list):
    # creating blank dataframe to append others to
    all_dataframes = []

    # Loop through each file, read it into a DataFrame, and append to the list
    for filename in csv_list:
        print(filename)

        file_path = os.path.join(setup.PROJECT_DIR, setup.RESULTS_FOLDER, filename)
        df = pd.read_csv(file_path)
        all_dataframes.append(df)

    # Concatenate all DataFrames into a single DataFrame
    combined_df = pd.concat(all_dataframes, ignore_index=True)
    combined_df['QC_file_duration'] = round((combined_df['QC_file_duration']/60/60), 2)

    return combined_df


# Save the combined DataFrame to a CSV file
def save_combined_csv(combined_df, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Created output folder: {output_folder}")

    output_file_path = os.path.join(output_folder, setup.PROJECT_NAME+'_combined_results.csv')
    combined_df.to_csv(output_file_path, index=False)
    print(f"Combined DataFrame saved to {output_file_path}")

# Creating log
def create_log(log_header):
    qc_log = docx.Document()
    qc_log.add_heading(log_header)
    return qc_log

# Save log
def save_log(log, output_file_path, date):
    log_path = os.path.join(output_file_path, f'QC_diagnostics_log_{NAME_JOB_LIST}_{date}.docx')
    log.save(log_path)


# --- Adding text to log --- #
def add_text(log, text_to_log, x, y, z):
    """
    This function adds text to the verification log
    :param log: The verification log that the output gets printed to.
    :param text_to_log: The explanatory text that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_to_log}")
    run.bold = True
    run.font.color.rgb = RGBColor(x, y, z)

# --- Adding text to log --- #
def add_description(log, description):
    """
    This function adds a description to the verification log
    :param log: The verification log that the output gets printed to.
    :param description: The explanatory text that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{description}")
    run.font.color.rgb = RGBColor(0, 0, 0)


# --- Adding PATT recommendation to log --- #
def patt_recommend(log, recommendation, x, y, z):
    """
    This function adds a PATT recommendation to the verification log
    :param log: The verification log that the output gets printed to.
    :param description: The PATT recommendation that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{recommendation}")
    run.font.color.rgb = RGBColor(x, y, z)

# --- Adding text to log if no errors --- #
def add_text_no_error(text_no_error, log):
    """
    This function adds text to the verification log
    :param log: The verification log that the output gets printed to.
    :param text_no_error: The explanatory text that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_no_error}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 155, 0)


# Check setup variables
def setup_var(log, df, var, comparison_operator, setup_var, text_to_log, extra_text, description, recommendation, x, y, z, list_headers, text_no_error):
    # Flagging if any files was not set up correctly
    flag = False
    flagged_files = set()
    grouped = df.groupby('id')

    for file_id, group in grouped:
        for _, row in group.iterrows():
                if operators[comparison_operator](row[var], setup_var):
                    flagged_files.add(file_id)
                    break

    # Adding table to log if any devices was not setup correctly
    if flagged_files:
        add_text(log, text_to_log, 0, 0, 0)

        if extra_text =='Yes':
            add_description(log, description)
            patt_recommend(log, recommendation, x, y, z)
        else:
            pass

        table = log.add_table(rows=1, cols=len(list_headers))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        headers = list_headers
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            run.bold = True

        # Printing information about the files that was flagged
        for file_id in flagged_files:
            group = grouped.get_group(file_id)

            for _, row in group.iterrows():
                row_cells = table.add_row().cells

                for idx, header in enumerate(list_headers):
                    if header in row:
                        row_cells[idx].text = str(row[header])
                    else:
                        row_cells[idx].text = ''
    else:
        add_text_no_error(text_no_error, log)

    log.add_paragraph("\n")
    return qc_log


if __name__ == '__main__':
    # Combining the QC_diagnostic output for all files into one csv
    csv_files = create_filelist(os.path.join(setup.PROJECT_DIR, setup.RESULTS_FOLDER))
    combined_df = read_and_append(csv_files)
    save_combined_csv(combined_df,  os.path.join(setup.PROJECT_DIR, setup.QC_OUTPUT))

    # Creating qc log with main qc variables to review
    qc_log = create_log("Main QC variables to review")

    # Checking if subject_code and filename matches:
    combined_df['subject_code'] = combined_df['subject_code'].apply(str)
    combined_df['id_flag'] = combined_df.apply(lambda row: int(row['subject_code'] not in row['file_filename']), axis=1)
    qc_log = setup_var(log=qc_log, df=combined_df, var='id_flag', comparison_operator='!=', setup_var=0, text_to_log='The subject_id and filename does not match for some of the files.', extra_text='Yes', description=' Check the files listed below:',
                       recommendation='PATT recommendations: Please check to find out who this data belongs to. Do the dates of wear match either of the IDs? \nFrom your own logs, did either of the IDs use the monitor that is listed as used? \nIf you cannot confirm who the data belongs to, this data will likely not be able to be used. Keep a note of the findings, so when it comes to processing the data, the team are not surprised when the ID within the file does not match what is stored in the filename.', x=255, y=0, z=0, list_headers=['subject_code', 'file_filename', 'device'], text_no_error='The subject code and filename matches for all files. No files to check.')

    # Checking if device number matched the device number in the filename (if OmGui is setup to add device number to filename:
    if device_filename == 'Yes':
        combined_df['device_flag'] = combined_df.apply(lambda row: int(str(row['device']) not in row['file_filename']), axis=1)
        qc_log = setup_var(log=qc_log, df=combined_df, var='device_flag', comparison_operator='!=', setup_var=0, text_to_log='The device number does not occur or do not match the device number in the filename for some files.', extra_text='Yes', description='Check the files listed below:',
                           recommendation='PATT recommendations: Check logs of who was given which monitor. Does the one in the filename or the one inside the header of the data match the one provided? \nIf the data is still on the device, double check by looking through OmGui. \nIf this is displayed incorrectly it would be recommended to set the monitor up for a few hours for a test, to see when downloaded, if it displays the correct device number.',
                           x=255, y=0, z=0, list_headers=['subject_code', 'file_filename', 'device'], text_no_error='The device number and filename matches for all files. No files to check.')

    # Checking if setup variables (frequency and range) matches:
    qc_log = setup_var(log=qc_log, df=combined_df, var='frequency', comparison_operator='!=', setup_var=frequency, text_to_log=f'Some devices were not set up with the frequency {frequency} hz:', extra_text='Yes', description='Check the files listed below:',
                       recommendation='PATT recommendations: Check your set up weblink (if that is being used) on all active computers/tablets to ensure that everyone setting devices up are using the correct link. Make sure staff are checking the settings before initialising the device. \nIf using OmGui make sure to reiterate to staff to check that the settings are correct before initialising the device.',
                       x=255, y=0, z=0, list_headers=['id', 'file_filename', 'device', 'frequency'], text_no_error=f'All devices were set up with the frequency {frequency} hz.')
    qc_log = setup_var(log=qc_log, df=combined_df, var='sample_range', comparison_operator='!=', setup_var=sample_range, text_to_log=f'Some devices were not set up with the sample range {sample_range} g:', extra_text='Yes', description='Check the files listed below:',
                       recommendation='PATT recommendations: Check your set up weblink (if that is being used) on all active computers/tablets to ensure that everyone setting devices up are using the correct link. Make sure staff are checking the settings before initialising the device. \nIf using OmGui make sure to reiterate to staff to check that the settings are correct before initialising the device.',
                       x=255, y=0, z=0, list_headers=['id', 'file_filename', 'device', 'sample_range'], text_no_error=f'All devices were set up with the range {sample_range} g.')

    # Calculating file duration in days and printing out QC timestamps
    combined_df['file_duration_days'] = round(combined_df['QC_file_duration']/24, 2)
    combined_df['duration_flag'] = ((combined_df['file_duration_days'] < min_file_duration) | (combined_df['file_duration_days'] > max_file_duration)).astype(int)

    qc_log = setup_var(log=qc_log, df=combined_df, var='duration_flag', comparison_operator='!=', setup_var=0, text_to_log=f'The duration of some files is outside the range {min_file_duration} - {max_file_duration} days.', extra_text='Yes',
                       description='Check the files listed below:', recommendation=f'PATT recommendations: If the file duration is below {min_file_duration} days, check the starting battery %, was this lower than expected? \nIf the battery started at >90%, it is recommended that you test that device out before continuing to use it. This could be the first sign of the battery beginning to fail. \nIf the monitor lasts for the desired length on a re-test, it can be continued to be re-used.',
                       x=255, y=0, z=0, list_headers=['id', 'QC_first_timestamp', 'QC_last_timestamp', 'file_duration_days', 'QC_first_battery_pct'], text_no_error=f'The file duration was between {min_file_duration} - {max_file_duration} days for all files. No files to check.')


    # Checking if there was any battery issues
    qc_log = setup_var(log=qc_log, df=combined_df, var='QC_check_battery', comparison_operator='!=', setup_var=False,
                       text_to_log='Some files had battery issues. Check the potential battery issues listed below for the files in the table:',
                       extra_text='Yes', description='All the checks below are displayed in one table, review to see which: \nWas device set up with less than 80% battery: QC_first_battery_pct < 80%. \nDid the battery die: QC_lowest_battery_pct ~0. \nDid the battery die faster than 25% a day: QC_max_discharge > 25.',
                       recommendation='PATT recommendations: If the device was set up with less than 80% battery, check if the device can be charged to ~100%. \nIf the lowest battery percentage was ~0 and/or the max discharge was above 25%, check if the first battery percentage was above 80 %. Set the device up to a 7 day test and see if the battery lasts and does not discharge more than 25% a day.', x=255, y=0, z=0,
                       list_headers=['id', 'QC_first_battery_pct', 'QC_lowest_battery_pct', 'QC_max_discharge'], text_no_error='There were no battery issues, no files to check.')

    # Checking if there was any anomalies
    # First generating a flag to indicate if any anomalies
    qc_log['anomaly_flag'] = ((combined_df['QC_Anomaly_A'] > 0) | (combined_df['QC_Anomaly_B'] > 0) |
                              (combined_df['QC_Anomaly_C'] > 0) | (combined_df['QC_Anomaly_D'] > 0) |
                              (combined_df['QC_Anomaly_E'] > 0) | (combined_df['QC_Anomaly_F'] > 0) |
                              (combined_df['QC_Anomaly_G'] > 0)).astype(int)

    qc_log = setup_var(log=qc_log, df=combined_df, var='anomaly_flag', comparison_operator='!=', setup_var=0, text_to_log='Some files had axis anomalies. Check the files listed below:',
                       extra_text='Yes', description='Each of the anomalies relates to timestamp issues where the device has lost its ability to keep track of time.', recommendation='PATT recommendations: It is recommended to remove the device from the pool and not reuse it. \nThe data should be cleaned automatically, but make sure to check through the results.', x=255, y=0, z=0,
                       list_headers=['id', 'device', 'QC_anomaly_A', 'QC_anomaly_B', 'QC_anomaly_C', 'QC_anomaly_D', 'QC_anomaly_E', 'QC_anomaly_F', 'QC_anomaly_G'], text_no_error='No anomalies were found. No files to check.')

    # Checking minimum and maximum values for all axis
    # First generating a flag to indicate if any axis falls outside +-1.2g
    combined_df['axis_flag'] = ((combined_df['QC_X_min'] < -1.2) | (combined_df['QC_X_max'] > 1.2) |
                                (combined_df['QC_Y_min'] < -1.2) | (combined_df['QC_Y_max'] > 1.2) |
                                (combined_df['QC_Z_min'] < -1.2) | (combined_df['QC_Z_max'] > 1.2)).astype(int)

    qc_log = setup_var(log=qc_log, df=combined_df, var='axis_flag', comparison_operator='!=', setup_var=0, text_to_log='The axis minimum and maximum values falls outside the range +-1.2G for some files.',
                       extra_text='Yes', description='Check the files below:', recommendation='PATT recommendations: For any devices displayed below with 1 or more axis falling outside the range: remove the monitor from the pool and do not reuse.', x=255, y=0, z=0,
                       list_headers=['id', 'device', 'QC_X_min', 'QC_X_max', 'QC_Y_min', 'QC_Y_max', 'QC_Z_min', 'QC_Z_max'], text_no_error='The axis minimum and maximum are within the expected range of +-1.2G for all files. No files to check.')


    # Checking wear duration and device was worn enough in each quadrant
    qc_log = setup_var(log=qc_log, df=combined_df, var='QC_total_hours_wear', comparison_operator='<', setup_var=min_wear, text_to_log=f'Some wear periods are shorter than the required {min_wear} hours. Check the files below:',
                       extra_text='Yes', description='Check how many wear periods that are too short: is it all files or a one off.', recommendation='PATT recommendations: Required wear duration is study dependent, no recommendations provided.', x=255, y=0, z=0,
                       list_headers=['id', 'device', 'QC_total_hours_wear', 'QC_file_duration'], text_no_error=f'All wear periods are above the required {min_wear} hours. No files to check')

    combined_df['quadrant_flag'] = ((combined_df['QC_hours_wear_quadrant_0'] < quad_wear) | (combined_df['QC_hours_wear_quadrant_1'] < quad_wear) | (combined_df['QC_hours_wear_quadrant_2'] < quad_wear) | (combined_df['QC_hours_wear_quadrant_3'] < quad_wear)).astype(int)

    qc_log = setup_var(log=qc_log, df=combined_df, var='quadrant_flag', comparison_operator='!=', setup_var=0, text_to_log=f'Some files have less than {quad_wear} hours wear for one or more of the quadrants. Check the files below:',
                       extra_text='Yes', description='Quadrant 0 = 00:00-05:59, Quadrant 1 = 06:00-11:59, Quadrant 2 = 12:00-17:59, Quadrant 3 = 18:00-23:59.', recommendation='PATT recommendations: Required wear duration in each quadrant is study dependent, no recommendations provided.', x=255, y=0, z=0,
                       list_headers=['id', 'QC_total_hours_wear', 'QC_hours_wear_quadrant_0', 'QC_hours_wear_quadrant_1', 'QC_hours_wear_quadrant_2', 'QC_hours_wear_quadrant_3'], text_no_error=f'All files have more than {quad_wear} hours wear in each quadrant. No files to check.')

    save_log(qc_log, os.path.join(setup.PROJECT_DIR, setup.QC_OUTPUT), DATE)
